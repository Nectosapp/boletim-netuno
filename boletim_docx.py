# -*- coding: utf-8 -*-
"""
Netuno — Boletim de Mercado — Word (.docx) Edition
Gera documento Word com o conteúdo do boletim diário e envia por email como anexo.
"""

import os
import sys
import time
import smtplib
import traceback
from datetime import datetime, timedelta, timezone
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import docx.opc.constants

# Importa funções de coleta do boletim principal
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from boletim_netuno_v2 import (
    collect_news, collect_radar, yahoo_quote,
    fp, fpct, translate_to_pt,
    RADAR_BR, RADAR_US, B3_TICKERS, GLOBAL_TICKERS, TZ_BR,
)

# ================== CONFIG ==================
DESTINATARIOS_DOCX = [
    "gustavoportugalhamer@gmail.com",
    "Byanca.tavelli@uranofin.com",
]
GMAIL_USER = os.environ.get("GMAIL_USER", "gustavoportugalhamer@gmail.com")
GMAIL_APP_PASSWORD = os.environ.get("GMAIL_APP_PASSWORD", "")
TEST_MODE = os.environ.get("BOLETIM_DOCX_TEST_MODE", "false").lower() == "true"
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output_docx")

# Cores Netuno
NAVY = RGBColor(0x1A, 0x27, 0x44)
TEAL = RGBColor(0x00, 0xAC, 0xAD)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x28, 0xA7, 0x45)
RED = RGBColor(0xDC, 0x35, 0x45)
GRAY = RGBColor(0x6C, 0x75, 0x7D)
DARK = RGBColor(0x33, 0x33, 0x33)
LINK_BLUE = RGBColor(0x00, 0x66, 0xCC)


# ================== HYPERLINK (python-docx não tem API nativa) ==================
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0066CC")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)
    run_el.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run_el.append(t)

    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


# ================== HELPERS ==================
def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_colored_text(cell, text, color, bold=False, size=9):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"


def format_table_header(table):
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "1A2744")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = "Calibri"


def pcolor(v):
    if v is None:
        return GRAY
    try:
        f = float(v)
        return GREEN if f > 0 else (RED if f < 0 else GRAY)
    except Exception:
        return GRAY


def format_data():
    return (
        datetime.now(TZ_BR)
        .strftime("%d de %B de %Y")
        .replace("January", "Janeiro").replace("February", "Fevereiro")
        .replace("March", "Março").replace("April", "Abril")
        .replace("May", "Maio").replace("June", "Junho")
        .replace("July", "Julho").replace("August", "Agosto")
        .replace("September", "Setembro").replace("October", "Outubro")
        .replace("November", "Novembro").replace("December", "Dezembro")
    )


# ================== DOCUMENT SECTIONS ==================
def configure_styles(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK
    for level, size in [(1, 18), (2, 13), (3, 11)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = NAVY
        h.font.bold = True


def create_header(doc):
    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GRUPO NETUNO")
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY
    run.font.bold = True
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Inteligência de Mercado")
    run2.font.size = Pt(10)
    run2.font.color.rgb = TEAL
    run2.font.bold = True
    run2.font.name = "Calibri"

    # Linha separadora
    doc.add_paragraph("_" * 70).runs[0].font.color.rgb = TEAL

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Boletim Diário de Mercado")
    run3.font.size = Pt(16)
    run3.font.color.rgb = NAVY
    run3.font.bold = True
    run3.font.name = "Calibri"

    hora = datetime.now(TZ_BR).strftime("%H:%M")
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f"{format_data()} • {hora} (Brasília)")
    run4.font.size = Pt(10)
    run4.font.color.rgb = GRAY
    run4.font.name = "Calibri"

    doc.add_paragraph()  # espaço


def add_destaques(doc, news):
    doc.add_heading("DESTAQUES DO DIA", level=2)

    cats = [
        ("🌎 Internacional", news.get("Internacional", []), True),
        ("🇧🇷 Brasil", news.get("Brasil", []), False),
        ("🏢 Empresas", news.get("Empresas", []), False),
    ]

    for title, items, traduzir in cats:
        if not items:
            continue
        p_cat = doc.add_paragraph()
        run_cat = p_cat.add_run(title)
        run_cat.font.size = Pt(11)
        run_cat.font.color.rgb = TEAL
        run_cat.font.bold = True
        run_cat.font.name = "Calibri"

        for it in items[:3]:
            titulo = (it.get("titulo") or "").strip()
            if traduzir:
                titulo = translate_to_pt(titulo)
            link = it.get("link", "")
            fonte = it.get("fonte") or it.get("srcname") or ""

            p = doc.add_paragraph(style="List Bullet")
            if link:
                add_hyperlink(p, titulo, link)
            else:
                run = p.add_run(titulo)
                run.font.size = Pt(10)
            run_src = p.add_run(f"  — {fonte}")
            run_src.font.size = Pt(8)
            run_src.font.color.rgb = GRAY
            run_src.font.name = "Calibri"

    doc.add_paragraph()


def generate_insights(quotes_gl):
    insights = []
    sp = quotes_gl.get("ES=F", {})
    sp_pct = sp.get("regularMarketChangePercent")
    if sp_pct is not None:
        try:
            v = float(sp_pct)
            if v > 0:
                insights.append(f"Futuros de S&P 500 operam em alta de {fpct(sp_pct)}, sinalizando abertura positiva nos EUA")
            elif v < 0:
                insights.append(f"Futuros de S&P 500 recuam {fpct(sp_pct)}, indicando cautela nos mercados globais")
        except Exception:
            pass

    brl = quotes_gl.get("BRL=X", {})
    brl_pct = brl.get("regularMarketChangePercent")
    brl_px = brl.get("regularMarketPrice")
    if brl_pct is not None and brl_px is not None:
        try:
            v = float(brl_pct)
            px = f"{float(brl_px):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            if v > 0:
                insights.append(f"Dólar avança {fpct(brl_pct)} e é negociado a R$ {px}")
            elif v < 0:
                insights.append(f"Dólar recua {fpct(brl_pct)} e é negociado a R$ {px}")
            else:
                insights.append(f"Dólar estável a R$ {px}")
        except Exception:
            pass

    oil = quotes_gl.get("CL=F", {})
    oil_pct = oil.get("regularMarketChangePercent")
    if oil_pct is not None:
        try:
            v = float(oil_pct)
            if abs(v) >= 0.5:
                d = "sobe" if v > 0 else "cai"
                insights.append(f"Petróleo WTI {d} {fpct(oil_pct)}, influenciando setor de commodities")
        except Exception:
            pass

    btc = quotes_gl.get("BTC-USD", {})
    btc_pct = btc.get("regularMarketChangePercent")
    btc_px = btc.get("regularMarketPrice")
    if btc_pct is not None and btc_px is not None:
        try:
            v = float(btc_pct)
            if abs(v) >= 1.0:
                px = f"US$ {float(btc_px):,.0f}".replace(",", ".")
                d = "avança" if v > 0 else "recua"
                insights.append(f"Bitcoin {d} {fpct(btc_pct)} e é negociado a {px}")
        except Exception:
            pass

    return insights


def add_insights(doc, quotes_gl):
    insights = generate_insights(quotes_gl)
    if not insights:
        return
    doc.add_heading("💡 INSIGHTS DO MERCADO", level=2)
    for i in insights:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(i)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    doc.add_paragraph()


def add_radar(doc, radar_br, radar_us):
    if not radar_br and not radar_us:
        return
    doc.add_heading("RADAR DE ABERTURA DE MERCADO", level=2)

    def radar_table(items, flag_title):
        if not items:
            return
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run(flag_title)
        run_sub.font.size = Pt(11)
        run_sub.font.color.rgb = TEAL
        run_sub.font.bold = True

        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Empresa"
        hdr[1].text = "Notícia"
        hdr[2].text = "Consenso"
        hdr[3].text = "Alvo"
        format_table_header(table)

        for it in items:
            row = table.add_row().cells
            add_colored_text(row[0], f"{it['nome']} ({it['ticker']})", NAVY, bold=True, size=9)
            # Notícia com link clicável
            row[1].text = ""
            p_news = row[1].paragraphs[0]
            if it.get("link"):
                add_hyperlink(p_news, it["titulo"][:80], it["link"])
            else:
                run_t = p_news.add_run(it["titulo"][:80])
                run_t.font.size = Pt(8)
            run_pub = p_news.add_run(f"\n{it.get('pub', '')}")
            run_pub.font.size = Pt(7)
            run_pub.font.color.rgb = GRAY
            run_pub.font.name = "Calibri"

            add_colored_text(row[2], it.get("consenso", "—"), NAVY, size=8)
            add_colored_text(row[3], it.get("alvo", "—"), DARK, size=8)

        # Ajusta larguras
        for row in table.rows:
            row.cells[0].width = Cm(3.5)
            row.cells[1].width = Cm(8)
            row.cells[2].width = Cm(2.5)
            row.cells[3].width = Cm(3)

    radar_table(radar_br, "🇧🇷 Brasil")
    radar_table(radar_us, "🇺🇸 EUA")
    doc.add_paragraph()


def add_cotacoes(doc, title, emoji, tickers, quotes, show_ticker=True):
    doc.add_heading(f"{emoji} {title}", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Ativo"
    hdr[1].text = "Ticker" if show_ticker else ""
    hdr[2].text = "Preço"
    hdr[3].text = "Variação"
    format_table_header(table)

    for sym, label in tickers.items():
        d = quotes.get(sym, {})
        px = d.get("regularMarketPrice")
        pct = d.get("regularMarketChangePercent")
        tk = sym.replace(".SA", "") if show_ticker else ""

        row = table.add_row().cells
        add_colored_text(row[0], label, DARK, size=9)
        add_colored_text(row[1], tk, GRAY, size=8)
        add_colored_text(row[2], fp(sym, px), DARK, bold=True, size=9)
        add_colored_text(row[3], fpct(pct), pcolor(pct), bold=True, size=9)
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Zebra striping
        idx = list(tickers.keys()).index(sym)
        if idx % 2 == 1:
            for cell in row:
                set_cell_shading(cell, "F8F9FA")

    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(2.5)
        row.cells[2].width = Cm(4)
        row.cells[3].width = Cm(3)

    doc.add_paragraph()


def add_footer(doc):
    doc.add_paragraph("_" * 70).runs[0].font.color.rgb = TEAL
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Grupo Netuno | R2F Capital\n"
        "Este boletim é produzido automaticamente com informações de fontes públicas confiáveis.\n"
        "As informações não constituem recomendação de investimento.\n\n"
        "Fontes: InfoMoney, Valor Investe, Exame, Bloomberg, Reuters, FT, Seeking Alpha, "
        "Brazil Journal, Yahoo Finance"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.font.name = "Calibri"
    run.font.italic = True


# ================== EMAIL COM ANEXO ==================
def enviar_email_com_anexo(dest_list, assunto, corpo_texto, docx_path):
    gmail_user = GMAIL_USER.strip()
    gmail_pass = GMAIL_APP_PASSWORD.strip().replace(" ", "")
    if not gmail_user or not gmail_pass:
        print("[ERRO] Credenciais Gmail não encontradas.")
        return

    msg = MIMEMultipart("mixed")
    msg["Subject"] = assunto
    msg["From"] = gmail_user
    msg["To"] = ", ".join(dest_list)
    msg.attach(MIMEText(corpo_texto, "plain", "utf-8"))

    with open(docx_path, "rb") as f:
        part = MIMEBase(
            "application",
            "vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header(
            "Content-Disposition",
            f'attachment; filename="{os.path.basename(docx_path)}"',
        )
        msg.attach(part)

    print("[*] Conectando ao SMTP Gmail...")
    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(gmail_user, gmail_pass)
        server.sendmail(gmail_user, dest_list, msg.as_string())
    print(f"[OK] Email com anexo enviado para {len(dest_list)} destinatário(s).")


# ================== MAIN ==================
def main():
    print("[*] Boletim DOCX — Iniciando...")

    print("[*] Coletando notícias...")
    news = collect_news()
    print(f"[+] BR={len(news['Brasil'])}, INTL={len(news['Internacional'])}, EMP={len(news['Empresas'])}")

    print("[*] Cotações...")
    quotes_br = yahoo_quote(list(B3_TICKERS.keys()))
    quotes_gl = yahoo_quote(list(GLOBAL_TICKERS.keys()))

    print("[*] Radar de Abertura...")
    radar_br, radar_us = [], []
    try:
        radar_br = collect_radar(RADAR_BR, False)
        radar_us = collect_radar(RADAR_US, True)
        print(f"[+] Radar: BR={len(radar_br)}, US={len(radar_us)}")
    except Exception as e:
        print(f"[WARN] Radar: {e}")

    # Gera documento Word
    print("[*] Gerando documento Word...")
    doc = Document()
    configure_styles(doc)
    create_header(doc)
    add_destaques(doc, news)
    add_insights(doc, quotes_gl)
    add_radar(doc, radar_br, radar_us)
    add_cotacoes(doc, "COTAÇÕES B3 BRASIL", "📈", B3_TICKERS, quotes_br, show_ticker=True)
    add_cotacoes(doc, "COTAÇÕES MERCADOS GLOBAIS", "🌐", GLOBAL_TICKERS, quotes_gl, show_ticker=False)
    add_footer(doc)

    # Salva
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    data_str = datetime.now(TZ_BR).strftime("%Y-%m-%d")
    filename = f"Boletim_Netuno_{data_str}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    print(f"[+] Documento salvo: {filepath}")

    # Envia por email
    if TEST_MODE:
        dest = ["gustavoportugalhamer@gmail.com"]
        print("[TEST MODE] Enviando apenas para teste.")
    else:
        dest = DESTINATARIOS_DOCX

    data_fmt = datetime.now(TZ_BR).strftime("%d/%m/%Y")
    assunto = f"[Grupo Netuno] Boletim Diário de Mercado — {data_fmt}"
    corpo = (
        f"Bom dia!\n\n"
        f"Segue em anexo o Boletim Diário de Mercado do Grupo Netuno ({data_fmt}).\n\n"
        f"Atenciosamente,\n"
        f"Grupo Netuno | R2F Capital"
    )
    enviar_email_com_anexo(dest, assunto, corpo, filepath)


if __name__ == "__main__":
    try:
        main()
        print("[OK] Boletim DOCX finalizado.")
    except Exception as e:
        print("[ERRO]", e)
        traceback.print_exc()
        exit(1)
